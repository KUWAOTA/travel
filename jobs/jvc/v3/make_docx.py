#!/usr/bin/env python3
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.replace("<br>", "\n"))
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    run.font.size = Pt(9)
    run.bold = bold


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 20, "1F3A5F"),
        ("Heading 2", 14, "1F3A5F"),
        ("Heading 3", 12, "1F2937"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8 if name != "Heading 1" else 0)
        style.paragraph_format.space_after = Pt(5)

    return doc


def add_run_text(paragraph, text, bold=False):
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    run.font.size = Pt(10.5)
    run.bold = bold
    return run


def parse_table(lines, idx):
    table_lines = []
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        table_lines.append(lines[idx].strip())
        idx += 1
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows, idx


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    rows = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_text(cell, text, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "E8EEF6")
    doc.add_paragraph()


def md_to_docx(md_path, out_path):
    doc = setup_document()
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    pending_para = []

    def flush_para():
        nonlocal pending_para
        if pending_para:
            p = doc.add_paragraph()
            add_run_text(p, "".join(pending_para))
            pending_para = []

    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw.strip()

        if not line:
            flush_para()
            i += 1
            continue

        if line.startswith("|"):
            flush_para()
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if line.startswith("# "):
            flush_para()
            p = doc.add_paragraph(style="Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run_text(p, line[2:], bold=True)
            i += 1
            continue

        if line.startswith("## "):
            flush_para()
            p = doc.add_paragraph(style="Heading 2")
            add_run_text(p, line[3:], bold=True)
            i += 1
            continue

        if line.startswith("### "):
            flush_para()
            p = doc.add_paragraph(style="Heading 3")
            add_run_text(p, line[4:], bold=True)
            i += 1
            continue

        if line.startswith("- "):
            flush_para()
            p = doc.add_paragraph(style="List Bullet")
            add_run_text(p, line[2:])
            i += 1
            continue

        if re.match(r"^[^:：]{1,18}[:：]\s*", line):
            flush_para()
            label, value = re.split(r"[:：]\s*", line, maxsplit=1)
            p = doc.add_paragraph()
            add_run_text(p, label + ": ", bold=True)
            add_run_text(p, value)
            i += 1
            continue

        if pending_para:
            pending_para.append("\n" + line)
        else:
            pending_para.append(line)
        i += 1

    flush_para()
    doc.save(out_path)


def main():
    OUT.mkdir(exist_ok=True)
    files = [
        ("職務経歴書_JVCラオス事業向け_v3.md", "職務経歴書_宇野拓磨_JVCラオス事業.docx"),
        ("応募動機作文_JVCラオス事業_v3.md", "応募動機作文_宇野拓磨_JVCラオス事業.docx"),
        ("JVCスタッフ申込書_記入案_v3.md", "JVCスタッフ申込書_記入案_宇野拓磨.docx"),
    ]
    for src, out in files:
        out_path = OUT / out
        md_to_docx(ROOT / src, out_path)
        print(out_path)


if __name__ == "__main__":
    main()
